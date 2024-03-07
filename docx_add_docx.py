from docx import Document
from docx.enum.section import WD_SECTION_START
from docxcompose.composer import Composer
#import tuozhuai
# 打开第一个 Word 文档,封面

#D:\周松霖\代码\封面\点之记封面.docx
#D:\周松霖\点之记\S205北流隆盛至平旦村二级公路点之记.docx
def docx_add_docx(list_word):

    result_document=Document()
    composer=Composer(result_document)
    for file in list_word:
        document=Document(file)
        
        composer.append(document)
    
        composer.save(r'D:\周松霖\点之记\S205北流隆盛至平旦村二级公路点之记+封面.docx')
        print('完成合并')
#document = Document(r'D:\周松霖\新建文件夹\result.docx')